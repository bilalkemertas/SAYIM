import streamlit as st
import pandas as pd
import datetime
import io
from zoneinfo import ZoneInfo
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from streamlit_gsheets import GSheetsConnection
import streamlit.components.v1 as components
from google import genai
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

TR_TZ = ZoneInfo("Europe/Istanbul")


def tr_now() -> datetime.datetime:
    """Türkiye yerel saatiyle şu anki zaman (tz bilgisi olmadan, sheet'e yazmak için)."""
    return datetime.datetime.now(TR_TZ).replace(tzinfo=None)


def tr_today() -> datetime.date:
    """Türkiye yerel tarihiyle bugün."""
    return tr_now().date()

# --- SAYFA AYARLARI (Mobil Uyumlu) ---
st.set_page_config(
    page_title="Satın Alma Görev Takip",
    page_icon="🛋️",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# --- KEEP-ALIVE: Streamlit'in inaktivite nedeniyle uykuya geçmesini engelle ---
# Streamlit Cloud, belirli bir süre trafik almayan uygulamaları uykuya alır.
# Aşağıdaki script arka planda periyodik olarak sunucuya sessiz bir istek
# göndererek uygulamayı "aktif" tutar. Sayfayı yenilemez, kullanıcıyı etkilemez.
KEEP_ALIVE_INTERVAL_MS = 4 * 60 * 1000  # 4 dakikada bir ping (5 dk sınırının altında kalır)
keep_alive_html = f"""
<script>
(function() {{
    setInterval(function() {{
        fetch(window.location.href, {{method: 'HEAD', cache: 'no-store'}}).catch(function(e) {{}});
    }}, {KEEP_ALIVE_INTERVAL_MS});
}})();
</script>
"""
components.html(keep_alive_html, height=0, width=0)

WORKSHEET = "tasks"
COLUMNS = ["id", "tarih", "gorev", "aciklama", "sorumlu", "durum",
           "silindi", "tamamlanma_tarihi", "silinme_tarihi", "silen",
           "created_at", "olusturan", "tamamlayan"]

# --- GOOGLE SHEETS BAĞLANTISI ---
@st.cache_resource
def init_connection():
    return st.connection("gsheets", type=GSheetsConnection)

try:
    conn = init_connection()
except Exception as e:
    st.error(f"Google Sheets bağlantı hatası: {e}")
    st.stop()


def load_tasks() -> pd.DataFrame:
    try:
        df = conn.read(worksheet=WORKSHEET, ttl=0)
    except Exception as e:
        st.error(f"Veri okunurken hata oluştu: {e}")
        st.stop()

    if df is None or df.empty:
        df = pd.DataFrame(columns=COLUMNS)

    for col in COLUMNS:
        if col not in df.columns:
            df[col] = ""

    df = df.dropna(how="all")

    df["silindi"] = df["silindi"].fillna(False).astype(bool)
    for col in ["gorev", "aciklama", "sorumlu", "durum", "tarih", "id",
                "tamamlanma_tarihi", "silinme_tarihi", "silen", "created_at",
                "olusturan", "tamamlayan"]:
        df[col] = df[col].fillna("").astype(str)

    return df.reset_index(drop=True)


def save_tasks(df: pd.DataFrame):
    conn.update(worksheet=WORKSHEET, data=df[COLUMNS])
    st.cache_data.clear()


# --- KULLANICI GİRİŞİ (Şifreli) ---
def check_login():
    if st.session_state.get("logged_in"):
        return True

    st.title("🔒 Giriş Yap")
    users = dict(st.secrets.get("users", {}))

    with st.form("login_form"):
        username = st.selectbox("Kullanıcı adı:", ["Seçiniz..."] + list(users.keys()))
        password = st.text_input("Şifre:", type="password")
        submit = st.form_submit_button("Giriş Yap")

    if submit:
        if username == "Seçiniz...":
            st.error("Lütfen kullanıcı adı seçiniz.")
        elif username in users and password == users[username]:
            st.session_state["logged_in"] = True
            st.session_state["current_user"] = username
            st.rerun()
        else:
            st.error("Kullanıcı adı veya şifre hatalı.")

    return False


if not check_login():
    st.stop()

selected_user = st.session_state["current_user"]
DURUM_ETIKET = {
    "acik": "Açık", "devam": "Devam Eden", "kapali": "Kapalı",
    "problem": "Problem", "aksiyon": "Aksiyon", "hedef": "Hedef", "yonetim": "Yönetim Desteği"
}
PEOPLE = list(dict(st.secrets.get("users", {})).keys())

# --- ÇIKIŞ BUTONU ---
top_col1, top_col2 = st.columns([4, 1])
with top_col1:
    st.title("🛋️ Satın Alma Görev Takip")
    st.caption(f"Yatak Üretim Fabrikası · Canlı İş Takibi · Giriş yapan: **{selected_user}**")
with top_col2:
    if st.button("Çıkış"):
        st.session_state["logged_in"] = False
        st.rerun()

st.divider()

# --- VERİLERİ YÜKLE ---
df = load_tasks()

# Silinmemiş görevleri filtrele
df_active = df[df["silindi"] == False] if "silindi" in df.columns else df

# --- İSTATİSTİKLER ---
acik_sayisi = len(df_active[df_active["durum"] == "acik"])
devam_sayisi = len(df_active[df_active["durum"] == "devam"])
kapali_sayisi = len(df_active[df_active["durum"] == "kapali"])

col1, col2, col3 = st.columns(3)
col1.metric("Açık", acik_sayisi)
col2.metric("Devam Eden", devam_sayisi)
col3.metric("Kapalı", kapali_sayisi)

# --- SEKMELER ---
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📋 Görev Listesi", "➕ Yeni Görev Ekle", "📊 Excel'den Toplu Yükle",
    "📈 Rapor", "📄 Haftalık Rapor (Word)"
])

# TAB 1: GÖREV LİSTESİ
with tab1:
    search_query = st.text_input("🔍 Görevlerde ara...", "")
    status_filter = st.radio("Filtrele:", ["Açık", "Devam Eden", "Kapalı", "Tümü"], horizontal=True)

    filtered_df = df_active.copy()

    if status_filter == "Açık":
        filtered_df = filtered_df[filtered_df["durum"] == "acik"]
    elif status_filter == "Devam Eden":
        filtered_df = filtered_df[filtered_df["durum"] == "devam"]
    elif status_filter == "Kapalı":
        filtered_df = filtered_df[filtered_df["durum"] == "kapali"]

    if search_query:
        filtered_df = filtered_df[
            filtered_df["gorev"].str.contains(search_query, case=False, na=False) |
            filtered_df["aciklama"].str.contains(search_query, case=False, na=False)
        ]

    st.write(f"**Toplam {len(filtered_df)} görev gösteriliyor:**")

    for idx, row in filtered_df.iterrows():
        with st.expander(f"**{row['gorev']}** ({row.get('sorumlu') or 'Atanmadı'})"):
            st.caption(f"Mevcut durum: **{DURUM_ETIKET.get(row['durum'], row['durum'])}**")
            if row.get("tamamlayan"):
                st.caption(f"Tamamlayan: {row.get('tamamlayan')} · {row.get('tamamlanma_tarihi', '-')}")

            with st.form(f"edit_form_{row['id']}"):
                mevcut_sorumlu_index = (PEOPLE.index(row["sorumlu"]) + 1) if row["sorumlu"] in PEOPLE else 0
                yeni_sorumlu = st.selectbox("Sorumlu", [""] + PEOPLE, index=mevcut_sorumlu_index, key=f"sorumlu_{row['id']}")

                try:
                    mevcut_tarih = datetime.datetime.strptime(row["tarih"], "%Y-%m-%d").date()
                except (ValueError, TypeError):
                    mevcut_tarih = tr_today()
                yeni_tarih = st.date_input("Termin Tarihi", value=mevcut_tarih, key=f"tarih_{row['id']}")

                yeni_aciklama = st.text_area("Açıklama", value=row.get("aciklama", ""), key=f"aciklama_{row['id']}")

                durum_secenekleri = ["acik", "devam", "kapali"]
                mevcut_durum_index = durum_secenekleri.index(row["durum"]) if row["durum"] in durum_secenekleri else 0
                yeni_durum = st.selectbox(
                    "Durum", durum_secenekleri, index=mevcut_durum_index,
                    format_func=lambda x: DURUM_ETIKET[x], key=f"durum_{row['id']}"
                )

                guncelle = st.form_submit_button("💾 Güncelle")

                if guncelle:
                    df.loc[df["id"] == row["id"], "sorumlu"] = yeni_sorumlu
                    df.loc[df["id"] == row["id"], "tarih"] = str(yeni_tarih)
                    df.loc[df["id"] == row["id"], "aciklama"] = yeni_aciklama
                    df.loc[df["id"] == row["id"], "durum"] = yeni_durum

                    if yeni_durum == "kapali" and row["durum"] != "kapali":
                        df.loc[df["id"] == row["id"], "tamamlanma_tarihi"] = str(tr_today())
                        df.loc[df["id"] == row["id"], "tamamlayan"] = selected_user
                    elif yeni_durum != "kapali":
                        df.loc[df["id"] == row["id"], "tamamlanma_tarihi"] = ""
                        df.loc[df["id"] == row["id"], "tamamlayan"] = ""

                    save_tasks(df)
                    st.toast("Görev güncellendi!")
                    st.rerun()

            if st.button("🗑️ Sil", key=f"del_{row['id']}"):
                df.loc[df["id"] == row["id"], "silindi"] = True
                df.loc[df["id"] == row["id"], "silinme_tarihi"] = str(tr_today())
                df.loc[df["id"] == row["id"], "silen"] = selected_user
                save_tasks(df)
                st.toast("Görev silindi!")
                st.rerun()

# TAB 2: YENİ GÖREV EKLE
with tab2:
    with st.form("new_task_form"):
        f_gorev = st.text_input("Görev Tanımı*")
        f_aciklama = st.text_area("Açıklama / Detay")
        f_sorumlu = st.selectbox("Sorumlu", [""] + PEOPLE)
        f_tarih = st.date_input("Görev Tarihi", tr_today())

        submitted = st.form_submit_button("💾 Kaydet")
        if submitted:
            if not f_gorev.strip():
                st.error("Lütfen görev tanımını boş bırakmayın!")
            else:
                new_row = pd.DataFrame([{
                    "id": f"task-{int(tr_now().timestamp())}",
                    "gorev": f_gorev,
                    "aciklama": f_aciklama,
                    "sorumlu": f_sorumlu,
                    "tarih": str(f_tarih),
                    "durum": "acik",
                    "silindi": False,
                    "tamamlanma_tarihi": "",
                    "silinme_tarihi": "",
                    "silen": "",
                    "created_at": str(tr_now()),
                    "olusturan": selected_user,
                    "tamamlayan": "",
                }])
                df = pd.concat([df, new_row], ignore_index=True)
                save_tasks(df)
                st.success("Yeni görev başarıyla eklendi!")
                st.rerun()

# TAB 3: EXCEL IMPORT
def _tr_normalize(s):
    s = str(s).strip()
    ceviri = str.maketrans({
        "İ": "i", "I": "i", "ı": "i",
        "Ö": "o", "ö": "o", "Ü": "u", "ü": "u",
        "Ş": "s", "ş": "s", "Ç": "c", "ç": "c", "Ğ": "g", "ğ": "g",
    })
    return s.translate(ceviri).lower()


def excel_kolon_bul(columns, adaylar):
    normalized = {_tr_normalize(c): c for c in columns}
    for aday in adaylar:
        key = _tr_normalize(aday)
        if key in normalized:
            return normalized[key]
    return None


# --- ŞABLON ÜRETİMİ: "Satinalma_Rapor_Veri_Sablonu.xlsx" ile birebir uyumlu ---
GECERLI_DURUMLAR = ["kapali", "devam", "problem", "aksiyon", "hedef", "yonetim"]

NAVY = "1F3864"
GRAY_ROW = "F4F4F4"
ORANGE_ROW = "FFF1E6"
NOTE_GRAY = "888888"


def generate_template_xlsx() -> bytes:
    wb = Workbook()

    # --- Sayfa 1: Talimatlar ---
    ws_t = wb.active
    ws_t.title = "Talimatlar"
    talimat_satirlari = [
        "  SATIN ALMA HAFTALIK FAALİYET RAPORU — VERİ ŞABLONU",
        "",
        "Bu dosyayı doldurup uygulamadaki \"📊 Excel'den Toplu Yükle\" sekmesinden yükleyin.",
        "",
        "1) \"Gorevler\" sekmesine o haftanın görevlerini girin",
        "Her satır bir görevi temsil eder. Kolonlar:",
        "    tarih  →  Görevin başladığı tarih (YYYY-AA-GG, örn. 2026-08-04)",
        "    gorev  →  Görevin kısa başlığı",
        "    aciklama  →  Görevle ilgili detay / durum notu",
        "    sorumlu  →  Görevden sorumlu kişi (örn. F.Sarı, A.Gözbaşı)",
        "    durum  →  kapali (tamamlandı) veya devam (sürüyor) — hücreye tıklayınca açılır listeden seçin",
        "    silindi  →  Görev iptal/silindiyse TRUE, aksi halde FALSE bırakın",
        "    tamamlanma_tarihi  →  Sadece durum=kapali olan görevlerde doldurun (YYYY-AA-GG)",
        "",
        "2) Problemler / Aksiyonlar / Hedefler / Yönetim Desteği notlarını EKLEMEK için",
        "Aynı sekmenin altına, görev satırları gibi yeni satırlar ekleyin ama:",
        "    durum kolonuna → problem  /  aksiyon  /  hedef  /  yonetim yazın (açılır listede bunlar da var)",
        "    gorev kolonuna → not metninin tamamını yazın",
        "    tarih, aciklama, sorumlu, tamamlanma_tarihi kolonlarını boş bırakabilirsiniz",
        "Bu satırlar görev sayılmaz; KPI hesaplarına dahil edilmez.",
        "",
        "3) Geçerli \"durum\" değerleri",
        "    kapali        → tamamlanmış görev",
        "    devam         → sürmekte olan görev",
        "    problem       → \"Karşılaşılan Problemler\" bölümüne not",
        "    aksiyon       → \"Alınan Aksiyonlar\" bölümüne not",
        "    hedef         → \"Bir Sonraki Haftanın Hedefleri\" bölümüne not",
        "    yonetim       → \"Yönetim Desteği Gerektiren Konular\" bölümüne not",
        "",
        "4) \"Gorevler\" sekmesindeki gri ve turuncu renkli satırlar örnektir",
        "Kendi verinizi girdikten sonra bu örnek satırları silebilir ya da üzerine yazabilirsiniz.",
        "",
        "5) Doldurduktan sonra",
        "Dosyayı .xlsx olarak kaydedin ve \"Excel'den Toplu Yükle\" sekmesinden yükleyip aktarın.",
    ]
    for i, satir in enumerate(talimat_satirlari, start=1):
        ws_t.cell(row=i, column=1, value=satir)
    ws_t.column_dimensions["A"].width = 100

    # --- Sayfa 2: Gorevler ---
    ws = wb.create_sheet("Gorevler")
    headers = ["tarih", "gorev", "aciklama", "sorumlu", "durum", "silindi", "tamamlanma_tarihi"]
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=NAVY)

    ornek_gorevler = [
        ("2026-08-04", "Örnek: Wind 11 kumaş alımı", "Bursa Dawis firmasından 8 mt kumaş temin edildi", "A.Gözbaşı", "kapali", "FALSE", "2026-08-05"),
        ("2026-08-06", "Örnek: Hidrant yangın dolabı teklif toplama", "3. Hangar için 28 adet; firmalardan dönüş bekleniyor", "F.Sarı", "devam", "FALSE", None),
    ]
    ornek_notlar = [
        (None, "Örnek: Freze bıçağı bileme fiyatı yüksek bulunarak iptal edildi", None, None, "problem", "FALSE", None),
        (None, "Örnek: Feryal Fermuar (2.000 EUR) fiyat farkı için mail gönderildi", None, None, "aksiyon", "FALSE", None),
        (None, "Örnek: Fiyat hatası çalışmalarının sonuçlandırılması", None, None, "hedef", "FALSE", None),
        (None, "Örnek: Hidrant yangın dolabı teklif onayının üst yönetimden verilmesi", None, None, "yonetim", "FALSE", None),
    ]

    r = 2
    for row_data in ornek_gorevler:
        for c, val in enumerate(row_data, start=1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.fill = PatternFill("solid", fgColor=GRAY_ROW)
            cell.font = Font(color="333333")
        r += 1
    for row_data in ornek_notlar:
        for c, val in enumerate(row_data, start=1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.fill = PatternFill("solid", fgColor=ORANGE_ROW)
            cell.font = Font(color="333333")
        r += 1

    r += 1  # boş satır
    note_cell = ws.cell(row=r, column=1, value=(
        "↑ Gri satırlar örnek GÖREV, turuncu satırlar örnek NOT satırıdır. "
        "Kendi verinizi 2. satırdan başlayarak girin / üzerine yazın. "
        "Ayrıntılı kurallar için \"Talimatlar\" sekmesine bakın."
    ))
    note_cell.font = Font(color=NOTE_GRAY, italic=True)
    r += 1

    # boş, doldurulmaya hazır satırlar
    for _ in range(25):
        ws.cell(row=r, column=6, value="FALSE")
        r += 1

    ws.column_dimensions["A"].width = 12
    ws.column_dimensions["B"].width = 42
    ws.column_dimensions["C"].width = 50
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["E"].width = 12
    ws.column_dimensions["F"].width = 10
    ws.column_dimensions["G"].width = 18
    ws.freeze_panes = "A2"

    dv_durum = DataValidation(type="list", formula1=f'"{",".join(GECERLI_DURUMLAR)}"', allow_blank=True)
    ws.add_data_validation(dv_durum)
    dv_durum.add("E2:E500")

    dv_silindi = DataValidation(type="list", formula1='"FALSE,TRUE"', allow_blank=True)
    ws.add_data_validation(dv_silindi)
    dv_silindi.add("F2:F500")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


with tab3:
    st.caption("Excel dosyanızda en az bir 'Görev' sütunu bulunmalı. Sütun adları büyük/küçük harfe duyarlı değildir.")

    st.download_button(
        "⬇️ Şablon İndir (.xlsx)",
        data=generate_template_xlsx(),
        file_name="Satinalma_Rapor_Veri_Sablonu.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    uploaded_file = st.file_uploader("Excel dosyasını yükleyin (.xlsx)", type=["xlsx", "xls"])
    if uploaded_file:
        try:
            excel_df = pd.read_excel(uploaded_file, sheet_name="Gorevler") if "Gorevler" in pd.ExcelFile(uploaded_file).sheet_names else pd.read_excel(uploaded_file)
            st.write("Yüklenecek Veri Önizlemesi:", excel_df.head())

            gorev_col = excel_kolon_bul(excel_df.columns, ["Görev Tanımı", "Görev", "Gorev", "Gorev Tanimi", "Görev Adı"])
            aciklama_col = excel_kolon_bul(excel_df.columns, ["Açıklama", "Aciklama", "Detay", "Açıklama/Detay"])
            sorumlu_col = excel_kolon_bul(excel_df.columns, ["Sorumlu"])
            tarih_col = excel_kolon_bul(excel_df.columns, ["Tarih", "Görev Tarihi", "Termin", "Termin Tarihi"])
            durum_col = excel_kolon_bul(excel_df.columns, ["Durum"])
            silindi_col = excel_kolon_bul(excel_df.columns, ["Silindi"])
            tamamlanma_col = excel_kolon_bul(excel_df.columns, ["Tamamlanma Tarihi", "Tamamlanma_Tarihi", "Tamamlanma"])

            if not gorev_col:
                st.error(
                    "Görev sütunu bulunamadı. Excel'deki sütun başlıklarınız: "
                    + ", ".join(str(c) for c in excel_df.columns)
                    + " — lütfen sütunlardan birinin adı 'Görev' veya 'Görev Tanımı' olsun."
                )
            else:
                st.success(f"Görev sütunu olarak **'{gorev_col}'** kullanılacak.")

                if st.button("📤 Verileri Veritabanına Aktar"):
                    new_records = []
                    atlanan = 0
                    for idx, row in excel_df.iterrows():
                        gorev = row.get(gorev_col)
                        if pd.isna(gorev) or not str(gorev).strip():
                            atlanan += 1
                            continue

                        ham_tarih = row.get(tarih_col) if tarih_col else None
                        try:
                            gorev_tarihi = pd.to_datetime(ham_tarih).date() if ham_tarih and not pd.isna(ham_tarih) else tr_today()
                        except Exception:
                            gorev_tarihi = tr_today()

                        ham_durum = str(row.get(durum_col, "")).strip().lower() if durum_col else ""
                        secili_durum = ham_durum if ham_durum in (["acik"] + GECERLI_DURUMLAR) else "acik"

                        ham_silindi = str(row.get(silindi_col, "")).strip().upper() if silindi_col else ""
                        secili_silindi = ham_silindi == "TRUE"

                        ham_tamamlanma = row.get(tamamlanma_col) if tamamlanma_col else None
                        try:
                            secili_tamamlanma = str(pd.to_datetime(ham_tamamlanma).date()) if ham_tamamlanma and not pd.isna(ham_tamamlanma) else ""
                        except Exception:
                            secili_tamamlanma = ""

                        new_records.append({
                            "id": f"excel-{int(tr_now().timestamp())}-{idx}",
                            "gorev": str(gorev).strip(),
                            "aciklama": str(row.get(aciklama_col, "")).strip() if aciklama_col else "",
                            "sorumlu": str(row.get(sorumlu_col, "")).strip() if sorumlu_col else "",
                            "tarih": str(gorev_tarihi),
                            "durum": secili_durum,
                            "silindi": secili_silindi,
                            "tamamlanma_tarihi": secili_tamamlanma,
                            "silinme_tarihi": "",
                            "silen": "",
                            "created_at": str(tr_now()),
                            "olusturan": selected_user,
                            "tamamlayan": "",
                        })

                    if new_records:
                        df = pd.concat([df, pd.DataFrame(new_records)], ignore_index=True)
                        save_tasks(df)
                        st.success(f"{len(new_records)} adet yeni görev Excel'den aktarıldı! ({atlanan} satır boş olduğu için atlandı)")
                        st.rerun()
                    else:
                        st.warning(f"Aktarılacak kayıt bulunamadı. {atlanan} satır boş/geçersiz olduğu için atlandı.")
        except Exception as e:
            st.error(f"Excel okunurken hata oluştu: {e}")

# TAB 4: RAPOR
with tab4:
    st.subheader("📈 Görev Raporu")

    rapor_kaynak = df.copy()
    rapor_kaynak["goruntu_durum"] = rapor_kaynak.apply(
        lambda r: "silindi" if r["silindi"] else r["durum"], axis=1
    )

    r1, r2 = st.columns(2)
    durum_secim = r1.multiselect(
        "Durum", ["acik", "devam", "kapali", "problem", "aksiyon", "hedef", "yonetim", "silindi"],
        default=["acik", "devam", "kapali"],
        format_func=lambda x: {
            "acik": "Açık", "devam": "Devam Eden", "kapali": "Kapalı", "silindi": "Silindi",
            "problem": "Problem", "aksiyon": "Aksiyon", "hedef": "Hedef", "yonetim": "Yönetim Desteği"
        }[x]
    )
    sorumlu_listesi = sorted([s for s in rapor_kaynak["sorumlu"].unique() if s])
    sorumlu_secim = r2.multiselect("Sorumlu", sorumlu_listesi, default=sorumlu_listesi)

    r3, r4 = st.columns(2)
    tamamlayan_listesi = sorted([s for s in rapor_kaynak["tamamlayan"].unique() if s])
    tamamlayan_secim = r3.multiselect("Tamamlayan", tamamlayan_listesi, default=tamamlayan_listesi)

    tarih_araligi = r4.date_input(
        "Görev tarihi aralığı",
        value=(tr_today() - datetime.timedelta(days=30), tr_today())
    )

    rapor_df = rapor_kaynak.copy()
    rapor_df = rapor_df[rapor_df["goruntu_durum"].isin(durum_secim)]
    if sorumlu_secim:
        rapor_df = rapor_df[rapor_df["sorumlu"].isin(sorumlu_secim) | (rapor_df["sorumlu"] == "")]
    if tamamlayan_secim:
        rapor_df = rapor_df[rapor_df["tamamlayan"].isin(tamamlayan_secim) | (rapor_df["tamamlayan"] == "")]

    if len(tarih_araligi) == 2:
        try:
            rapor_df["_tarih_dt"] = pd.to_datetime(rapor_df["tarih"], errors="coerce")
            baslangic, bitis = tarih_araligi
            rapor_df = rapor_df[
                (rapor_df["_tarih_dt"].dt.date >= baslangic) & (rapor_df["_tarih_dt"].dt.date <= bitis)
                | (rapor_df["_tarih_dt"].isna())
            ]
            rapor_df = rapor_df.drop(columns=["_tarih_dt"])
        except Exception:
            pass

    st.write(f"**{len(rapor_df)} görev bulundu**")
    goruntu_df = rapor_df[["gorev", "sorumlu", "goruntu_durum", "tarih", "tamamlanma_tarihi",
                            "tamamlayan", "olusturan", "silinme_tarihi", "silen"]].rename(
        columns={"goruntu_durum": "durum"}
    )
    st.dataframe(goruntu_df, use_container_width=True, hide_index=True)

    xlsx_buf = io.BytesIO()
    with pd.ExcelWriter(xlsx_buf, engine="openpyxl") as writer:
        goruntu_df.to_excel(writer, index=False, sheet_name="Gorev Raporu")
        ws_rapor = writer.sheets["Gorev Raporu"]
        for c, kolon in enumerate(goruntu_df.columns, start=1):
            cell = ws_rapor.cell(row=1, column=c)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=NAVY)
            genislik = max(14, min(50, int(goruntu_df[kolon].astype(str).str.len().max() or 10) + 4))
            ws_rapor.column_dimensions[get_column_letter(c)].width = genislik
        ws_rapor.freeze_panes = "A2"
    xlsx_buf.seek(0)

    st.download_button(
        "⬇️ Excel (.xlsx) olarak indir",
        data=xlsx_buf.getvalue(),
        file_name="gorev_raporu.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# TAB 5: HAFTALIK RAPOR (WORD)

# --- YAPAY ZEKA İLE YÖNETİCİ ÖZETİ ---
def generate_ai_summary(donem_baslangic, donem_bitis, tamamlanan_donem, yeni_donem,
                         problem_donem, aksiyon_donem, hedef_donem, yonetim_donem,
                         acik_sayisi, devam_sayisi, kapali_sayisi):
    api_key = st.secrets.get("gemini_api_key")
    if not api_key:
        return None, "secrets.toml dosyasına 'gemini_api_key' eklenmemiş."

    def liste_metni(dataframe):
        if dataframe.empty:
            return "yok"
        return "\n".join(f"- {str(r.get('gorev','')).strip()}" for _, r in dataframe.iterrows())

    veri_metni = f"""Dönem: {donem_baslangic.strftime('%d.%m.%Y')} - {donem_bitis.strftime('%d.%m.%Y')}
Açık görev: {acik_sayisi}, Devam eden: {devam_sayisi}, Kapalı: {kapali_sayisi}

Bu dönem tamamlanan görevler:
{liste_metni(tamamlanan_donem)}

Bu dönem eklenen yeni görevler:
{liste_metni(yeni_donem)}

Karşılaşılan problemler:
{liste_metni(problem_donem)}

Alınan aksiyonlar:
{liste_metni(aksiyon_donem)}

Bir sonraki haftanın hedefleri:
{liste_metni(hedef_donem)}

Yönetim desteği gereken konular:
{liste_metni(yonetim_donem)}
"""

    try:
        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(
            model="gemini-3.5-flash-lite",
            contents=(
                "Aşağıdaki satın alma/depo görev verilerine dayanarak, üst yönetime sunulacak "
                "3-5 cümlelik kısa ve profesyonel bir Türkçe yönetici özeti yaz. "
                "Sadece özet metnini yaz; başlık, madde işareti veya ek açıklama ekleme.\n\n"
                f"{veri_metni}"
            )
        )
        return response.text.strip(), None
    except Exception as e:
        return None, str(e)


def add_report_table(doc, dataframe, columns, headers, empty_text="Bu dönemde kayıt yok."):
    if dataframe.empty:
        p = doc.add_paragraph(empty_text)
        p.italic = True
        return

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = str(row.get(col, "") or "-")


def add_notes_list(doc, dataframe, empty_text="Bu dönemde kayıt yok."):
    if dataframe.empty:
        p = doc.add_paragraph(empty_text)
        p.italic = True
        return

    for _, row in dataframe.iterrows():
        metin = str(row.get("gorev", "")).strip() or "-"
        sorumlu = str(row.get("sorumlu", "")).strip()
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(metin)
        if sorumlu:
            r = p.add_run(f"  ({sorumlu})")
            r.italic = True


def generate_weekly_report_docx(donem_baslangic, donem_bitis, tamamlanan_donem, yeni_donem,
                                 problem_donem, aksiyon_donem, hedef_donem, yonetim_donem,
                                 acik_sayisi, devam_sayisi, kapali_sayisi, hazirlayan,
                                 ai_ozet=None):
    doc = Document()

    title = doc.add_heading("Haftalık Faaliyet Raporu", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Dönem: ").bold = True
    meta.add_run(f"{donem_baslangic.strftime('%d.%m.%Y')} — {donem_bitis.strftime('%d.%m.%Y')}\n")
    meta.add_run("Hazırlayan: ").bold = True
    meta.add_run(f"{hazirlayan}\n")
    meta.add_run("Oluşturma Tarihi: ").bold = True
    meta.add_run(tr_today().strftime('%d.%m.%Y'))

    if ai_ozet:
        doc.add_heading("Yönetici Özeti", level=2)
        ozet_p = doc.add_paragraph(ai_ozet)
        ozet_p.italic = True

    doc.add_heading("Genel Durum (Anlık)", level=2)
    ozet = doc.add_paragraph()
    ozet.add_run(f"Açık görev sayısı: ").bold = True
    ozet.add_run(f"{acik_sayisi}\n")
    ozet.add_run(f"Devam eden görev sayısı: ").bold = True
    ozet.add_run(f"{devam_sayisi}\n")
    ozet.add_run(f"Kapalı görev sayısı: ").bold = True
    ozet.add_run(f"{kapali_sayisi}")

    doc.add_heading(f"Bu Dönemde Tamamlanan Görevler ({len(tamamlanan_donem)})", level=2)
    add_report_table(
        doc, tamamlanan_donem,
        columns=["gorev", "sorumlu", "tamamlanma_tarihi", "tamamlayan"],
        headers=["Görev", "Sorumlu", "Tamamlanma Tarihi", "Tamamlayan"]
    )

    doc.add_heading(f"Bu Dönemde Eklenen Yeni Görevler ({len(yeni_donem)})", level=2)
    add_report_table(
        doc, yeni_donem,
        columns=["gorev", "sorumlu", "durum", "tarih", "olusturan"],
        headers=["Görev", "Sorumlu", "Durum", "Tarih", "Ekleyen"]
    )

    doc.add_heading(f"Karşılaşılan Problemler ({len(problem_donem)})", level=2)
    add_notes_list(doc, problem_donem)

    doc.add_heading(f"Alınan Aksiyonlar ({len(aksiyon_donem)})", level=2)
    add_notes_list(doc, aksiyon_donem)

    doc.add_heading(f"Bir Sonraki Haftanın Hedefleri ({len(hedef_donem)})", level=2)
    add_notes_list(doc, hedef_donem)

    doc.add_heading(f"Yönetim Desteği Gerektiren Konular ({len(yonetim_donem)})", level=2)
    add_notes_list(doc, yonetim_donem)

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Bu rapor Satın Alma Görev Takip sisteminden {hazirlayan} tarafından oluşturulmuştur."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


with tab5:
    st.subheader("📄 Haftalık Faaliyet Raporu (Word)")
    st.caption("Raporu oluşturup Word (.docx) olarak indirin, dilediğiniz gibi düzenleyip kendi e-posta istemcinizden gönderin.")

    donem_col1, donem_col2 = st.columns(2)
    varsayilan_baslangic = tr_today() - datetime.timedelta(days=7)
    donem_baslangic = donem_col1.date_input("Dönem başlangıcı", varsayilan_baslangic, key="donem_baslangic")
    donem_bitis = donem_col2.date_input("Dönem bitişi", tr_today(), key="donem_bitis")

    ai_ozet_kullan = st.checkbox("🤖 Yapay zeka ile yönetici özeti ekle (raporun başına)")

    if st.button("📝 Raporu Oluştur"):
        if donem_baslangic > donem_bitis:
            st.error("Dönem başlangıcı, bitiş tarihinden sonra olamaz.")
        else:
            df_all = df_active.copy()
            df_all["_tarih_dt"] = pd.to_datetime(df_all["tarih"], errors="coerce")
            df_all["_tamamlanma_dt"] = pd.to_datetime(df_all["tamamlanma_tarihi"], errors="coerce")
            df_all["_created_dt"] = pd.to_datetime(df_all["created_at"], errors="coerce")

            tamamlanan_donem = df_all[
                (df_all["_tamamlanma_dt"].dt.date >= donem_baslangic) &
                (df_all["_tamamlanma_dt"].dt.date <= donem_bitis)
            ]
            yeni_donem = df_all[
                (df_all["_tarih_dt"].dt.date >= donem_baslangic) &
                (df_all["_tarih_dt"].dt.date <= donem_bitis) &
                (df_all["durum"].isin(["acik", "devam", "kapali"]))
            ]

            # Şablondaki "Problemler / Aksiyonlar / Hedefler / Yönetim Desteği" not satırları:
            # tarih genelde boş bırakıldığı için eklenme tarihine (created_at) göre dönem filtresi uygulanır.
            notlar_donem = df_all[
                (df_all["_created_dt"].dt.date >= donem_baslangic) &
                (df_all["_created_dt"].dt.date <= donem_bitis)
            ]
            problem_donem = notlar_donem[notlar_donem["durum"] == "problem"]
            aksiyon_donem = notlar_donem[notlar_donem["durum"] == "aksiyon"]
            hedef_donem = notlar_donem[notlar_donem["durum"] == "hedef"]
            yonetim_donem = notlar_donem[notlar_donem["durum"] == "yonetim"]

            ai_ozet_metni = None
            if ai_ozet_kullan:
                with st.spinner("Yapay zeka özeti hazırlanıyor..."):
                    ai_ozet_metni, ai_hata = generate_ai_summary(
                        donem_baslangic, donem_bitis, tamamlanan_donem, yeni_donem,
                        problem_donem, aksiyon_donem, hedef_donem, yonetim_donem,
                        acik_sayisi, devam_sayisi, kapali_sayisi
                    )
                if ai_hata:
                    st.warning(f"Yapay zeka özeti oluşturulamadı, rapor özetsiz devam ediyor: {ai_hata}")

            docx_buf = generate_weekly_report_docx(
                donem_baslangic, donem_bitis, tamamlanan_donem, yeni_donem,
                problem_donem, aksiyon_donem, hedef_donem, yonetim_donem,
                acik_sayisi, devam_sayisi, kapali_sayisi, selected_user,
                ai_ozet=ai_ozet_metni
            )

            st.session_state["rapor_docx"] = docx_buf.getvalue()
            st.session_state["rapor_dosya_adi"] = f"haftalik_rapor_{donem_baslangic}_{donem_bitis}.docx"
            st.success("Rapor oluşturuldu, aşağıdan indirebilirsiniz.")

    if "rapor_docx" in st.session_state:
        st.download_button(
            "⬇️ Word Raporunu İndir",
            data=st.session_state["rapor_docx"],
            file_name=st.session_state["rapor_dosya_adi"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- İMZA ---
st.markdown(
    """
    <style>
    .imza {
        position: fixed;
        bottom: 8px;
        left: 14px;
        font-size: 11px;
        color: #9a9a9a;
        opacity: 0.7;
        z-index: 100;
    }
    </style>
    <div class="imza">created by Bilal KEMERTAŞ 🛠️</div>
    """,
    unsafe_allow_html=True
)
